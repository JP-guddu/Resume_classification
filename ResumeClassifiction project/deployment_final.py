# -*- coding: utf-8 -*-
"""
Created on Thu Jun 22 23:35:20 2023

@author: sachin
"""

import os
import streamlit as st
import pickle
import re
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.ensemble import RandomForestClassifier
rf = RandomForestClassifier() 
import hydralit_components as hc
import pandas as pd
import time
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
st.set_option('deprecation.showPyplotGlobalUse', False)
from collections import Counter
from wordcloud import WordCloud
from sklearn.feature_extraction.text import CountVectorizer
from time import sleep

vectorizer = TfidfVectorizer(max_features=3757)
pickle_in = open('E:\project2\Resume1.pkl', 'rb')
classifier = pickle.load(pickle_in)

pickle_in1 = open("E:\project2\Tdfi1.pkl","rb")
vectorizer1 = pickle.load(pickle_in1)

def classify_resume(text):
    transformed_text = vectorizer1.transform([text])
    prediction = classifier.predict(transformed_text)
    if prediction[0] == 0:
        return "Peoplesoft"
    elif prediction[0] == 1:
        return "React Developer"
    elif prediction[0] == 2:
        return "SQL Developer"
    elif prediction[0] == 3:
        return "Workday"
    else:
        return "Unknown"
import os #Used for path
import re #Regex
import csv #text into csv file
from docx import Document #.Docx Convert into text
from PyPDF2 import PdfReader #pdf Convert into text
import win32com.client as win32 #used for open Word file Component Object Model (COM)
from win32com.client import constants
import pandas as pd
import glob
import streamlit as st
import spacy 
nlp = spacy.load('en_core_web_md') #loading spacy pre built model
from spacy.matcher import Matcher 
from docx2pdf import convert
from io import BytesIO
import win32com.client as win32
from docx import Document



@st.cache_data
def extract_text_from_pdf(file_path): #This function is used for extract text from pdf
    pdf_reader = PdfReader(file_path)
    #num_pages = len(pdf_reader.pages)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text()
    return text


    #for page in range(file_path):
       # page_obj = pdf_reader.getPage(page)
        #text += page_obj.extractText()

   # return text
    
@st.cache_data
def extract_text_from_docx(file_path): #This function is used for extract text from Docx
    doc = Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs]
    return '\n'.join(paragraphs)

@st.cache_data
def save_as_docx(file_contents): #If Doc is not readed that this function convert Doc into Docx.by open Docx file and saveAS into Doc
    # Opening MS Word
    file_extension = os.path.splitext(file_contents.name)[1]

    if file_extension.lower() == '.docx':
        return extract_text_from_docx(file_contents)
    elif file_extension.lower() == '.doc':
        pdf_path = os.path.splitext(file_contents.name)[0] + '.pdf'
        convert(file_contents, pdf_path)

        docx_path = os.path.splitext(file_contents.name)[0] + '.docx'
        doc = Document()
        doc.add_paragraph(extract_text_from_pdf(pdf_path))
        doc.save(docx_path)

        os.remove(pdf_path)  # Remove the temporary PDF file
        return extract_text_from_docx(docx_path)
import textract
import tempfile
@st.cache_data


def extract_text_from_doc(file_path):
    # Open Word application
    word = win32.gencache.EnsureDispatch('Word.Application')
    doc = word.Documents.Open(file_path)
    doc.Activate()

    # Create a new file path for DOCX
    new_file_abs = os.path.splitext(file_path)[0] + '.docx'

    # Save as DOCX format
    word.ActiveDocument.SaveAs(new_file_abs, FileFormat=win32.constants.wdFormatXMLDocument)
    doc.Close(False)

    # Extract text from the converted DOCX file
    docx_doc = Document(new_file_abs)
    paragraphs = [p.text for p in docx_doc.paragraphs]
    text = '\n'.join(paragraphs)

    # Remove the temporary DOCX file
    os.remove(new_file_abs)

    return text


@st.cache_data
def clean_text(text):
    cleaned_text = re.sub(r'\s+', ' ', text)  # Remove extra space
    return cleaned_text


@st.cache_data
def extract_skills(resume_text):
    nlp_text = nlp(resume_text)

    # removing stop words and implementing word tokenization
    tokens = [token.text for token in nlp_text if not token.is_stop]
    # Set of skills if skills match with this text than they append in skills column
    skills = ["c","my sql","nodejs","node.js","reactjs","javascript","html","css","javascript","angular js",
              "js","fcsm","sql developer","peopleSoft","mysql","sql","plsql","nosql","rdbms","ddl" ,"dml","dcl"
              ,"sql developer","core hcm","xml", "xslt", "eib", "core connectors","workday",'hcm',"peoplesoft admin",
               "PeopleSoft Admin","dba","" ]
    
    skillset = []
    
    # check for one-grams (example: python)
    for token in tokens:
        if token.lower() in skills:
            skillset.append(token)
    
    #check for bi-grams and tri-grams (example: machine learning)
    for token in nlp_text.noun_chunks:
        token = token.text.lower().strip()
        if token in skills:
            skillset.append(token)
    
  
    return ', '.join(set([i.capitalize() for i in skillset]))



# In name extraction we use POS tagging. 
matcher = Matcher(nlp.vocab) 
@st.cache_data
def extract_name(resume_text):
    nlp_text = nlp(resume_text)
    
    # First name and Last name are always Proper Nouns
    pattern = [{'POS': 'PROPN'}, {'POS': 'PROPN'}]
    
    matcher.add('NAME', [pattern], on_match = None)
    
    matches = matcher(nlp_text)
    
    for match_id, start, end in matches:
        span = nlp_text[start:end]
        return span.text


nlp = spacy.load('en_core_web_lg')
@st.cache_data
def extract_experiences(text): #Extract Experience from resume
    doc = nlp(text)

    experiences = []
    for sent in doc.sents:
        if 'experience' in sent.text.lower():
            experiences.append(sent.text)

    return ', '.join(set(i for i in experiences))



@st.cache_data
def ed(text):  #Extract Education from Resume
    doc = nlp(text)

    edu = []
    for sent in doc.sents:
        if ("educational qualification" and 'education') in sent.text.lower():
            edu.append(sent.text)
    return ', '.join(set(i for i in edu))



import re
@st.cache_data
def extract_website_links(text):
    # Regular expression pattern to match website links
    pattern = r'(https?://\S+)'

    # Find all matches using the pattern
    matches = re.findall(pattern, text)

    return ", ".join(matches)


@st.cache_data
def extract_phone_numbers(text):
    # Regular expression pattern to match phone numbers
    pattern = r'(?<!\d)(?:(?:\d{2}[-\s]?\d{8})|(?:\d{4}[-\s]?\d{6}))(?!\d)'

    # Find all matches using the pattern
    matches = re.findall(pattern, text)

    return ", ".join(matches)


@st.cache_data
def expDetails(Text):  #Extract experience in year from Resumes
    global sent
   
    Text = Text.split()
   
    for i in range(len(Text)-2):
        Text[i].lower()
        
        if Text[i] ==  'years':
            sent =  Text[i-2] + ' ' + Text[i-1] +' ' + Text[i] +' '+ Text[i+1] +' ' + Text[i+2]
            l = re.findall('\d*\.?\d+',sent)
            for i in l:
                a = float(i)
            return(a)
            return (sent)


def cleanResume(resumeText):
    if pd.isnull(resumeText) or not isinstance(resumeText, str):
        return ''
    #resumeText = re.sub(r'\d+', '', resumeText)
    resumeText = re.sub('•', '',resumeText)
    resumeText = re.sub(r'\s+', ' ', resumeText).strip()
    resumeText = re.sub('http\S+\s*', ' ', resumeText)  # remove URLs
    resumeText = re.sub('RT|cc', ' ', resumeText)  # remove RT and cc
    resumeText = re.sub('#\S+', '', resumeText)  # remove hashtags
    resumeText = re.sub('@\S+', '  ', resumeText)  # remove mentions
    resumeText = re.sub('[%s]' % re.escape("""!"#$%&'()*+,-./:;<=>?@[\]^_`{|}~"""), ' ', resumeText)  # remove punctuations
    resumeText = re.sub(r'[^\x00-\x7f]',r' ', resumeText) 
    resumeText = re.sub('\s+', ' ', resumeText)  # remove extra whitespace
    return resumeText        

#input_directory = r'C:\Users\abc\Desktop\Resume Classification\Resume1'
#output_csv = 'Resume_text.csv'

#with open(output_csv, 'w', newline='', encoding='utf-8') as csv_file:
    #writer = csv.writer(csv_file)
    #writer.writerow(['Text',"Resume"])  # Write header row
    #for file_path in glob.glob(os.path.join(input_directory, '**', '*'), recursive=True):
        #if not os.path.isfile(file_path):
            #continue

#make it look nice from the start
st.set_page_config(layout='wide',initial_sidebar_state='collapsed')
import streamlit as st




st.title('RESUME CLASSIFICATION')

# specify the primary menu definition
menu_data = [
    
    {'icon': "far fa-file-word", 'label':"Resume Classification"},
    {'icon': "far fa-chart-bar", 'label':"Data Analysis"},
    {'icon': "far fa-sticky-note", 'label':"About"},#no tooltip message
] 
over_theme = {'txc_inactive': '#FFFFFF'}
menu_id = hc.nav_bar(
    menu_definition=menu_data,
    override_theme=over_theme,
    login_name=None,
    hide_streamlit_markers=False, #will show the st hamburger as well as the navbar now!
    sticky_nav=True, #at the top or not
    sticky_mode='pinned', #jumpy or not-jumpy, but sticky or pinned
)
if menu_id == 'Resume Classification':

    def add_bg_image():
        st.markdown(
          f"""
          <style>
         .stApp {{
             background-image: url("http://7-themes.com/data_images/collection/8/4498058-white-backgrounds.jpg");
             background-attachment: fixed;
             background-size: cover
         }}
         </style>
         """,
         unsafe_allow_html=True)

    add_bg_image()

    with hc.HyLoader('Please Wait!',hc.Loaders.standard_loaders,index=5):
        time.sleep(0.8)

   # st.title("RESUME CLASSIFICATION")
        
    st.subheader('Upload Resumes')

    st.write(r'Note: Classifies only Peoplesoft, Workday, SQL Developer and ReactJS Developer Resumes')
    tab1, tab2 = st.tabs(["💾 Single File","📁 Multiple Files"])
    with tab1:

         file_upload = st.file_uploader("Load your Resume", type = [".pdf",".Doc",".Docx"])

    st.write('*Note: For different Resumes Results Reupload') 
    if file_upload is not None:
        path = file_upload.read()
        file_extenstion = file_upload.name.split(".")[-1]
        if file_extenstion == "pdf":
            text  = extract_text_from_pdf(file_upload)
            text1  = clean_text(text)
            text1 = cleanResume(text1)
            prediction = classify_resume(text1)
            #st.header("Designation")
            #st.write(prediction)

            st.header("The "+ file_upload.name +" is Applied for"+ " " + prediction + " " + "Profile")
            expander = st.expander("See Resume")
            expander.write(text)
            st.header("Designation")
            st.write(prediction)
            skills = extract_skills(text1)
            # Display extracted skills
            st.header("Skills:")
    
            st.text(skills)
            Name = extract_name(text1)
            st.header("Name:")
            st.text(Name)
            Experiences = extract_experiences(text1)
            st.header("Previous Experiences:")
            st.text(Experiences)
            Year = expDetails(text1)
            st.header("Experience:")
            st.text(Year)
            web = extract_website_links(text1)
            st.header("Links:")
            st.text(web)
            phone = extract_phone_numbers(text1)
            st.header("Contect Number:")
            st.text(phone)
       #if predicted == 'Workday':
          #  st.image("https://www.workday.com/content/dam/web/en-us/images/social/workday-og-theme.png",width=480)
       # elif predicted == 'SQL Developer':
          #  st.image("https://wallpaperaccess.com/full/2138094.jpg",width=480)
       # elif predicted == 'React Developer':
           # st.image("https://i0.wp.com/www.electrumitsolutions.com/wp-content/uploads/2020/12/wp4923992-react-js-wallpapers.png",width=480)
       # elif predicted == 'Peoplesoft':
           # st.image("https://s3.amazonaws.com/questoracle-staging/wordpress/uploads/2019/07/25164143/PeopleSoft-Now.jpg",width=480)
        elif file_extenstion == "docx":
            text = extract_text_from_docx(file_upload)
            text1  = clean_text(text)
            text1 = cleanResume(text1)
            prediction = classify_resume(text1)
            #st.header("Designation")
            st.header("The "+ file_upload.name +" is Applied for"+ " " + prediction + " " + "Profile")
            expander = st.expander("See Resume")
            expander.write(text)
            st.header("Designation")
            st.write(prediction)
            skills = extract_skills(text1)
            # Display extracted skills
            st.header("Skills:")
    
            st.text(skills)
            Name = extract_name(text1)
            st.header("Name:")
            st.text(Name)
            Experiences = extract_experiences(text1)
            st.header("Previous Experiences:")
            st.text(Experiences)
            Year = expDetails(text1)
            st.header("Experience:")
            st.text(Year)
            web = extract_website_links(text1)
            st.header("Links:")
            st.text(web)
            phone = extract_phone_numbers(text1)
            st.header("Contect Number:")
            st.text(phone)
    
    with tab2:
        st.write('Upload Folder Containing Multiple .docx,.pdf and .doc Files')

        file_type=pd.DataFrame(columns=['Uploaded File', 'Experience', 'Skills', 'Predicted Profile'], dtype=object)
        filename = []
        predicted = []
        experience = []
        skills = []

        upload_file2 = st.file_uploader('', type= ['docx',"Pdf","Doc"], accept_multiple_files=True)
        
        for doc_file in upload_file2:
            if doc_file is not None:
                if file_upload:
                    path = file_upload.read()
                    file_extenstion = file_upload.name.split(".")[-1]
    
                    if file_extenstion == "pdf":
                        filename.append(doc_file.name)   
                        cleaned=clean_text(extract_text_from_pdf(doc_file))
                        cleaned = cleanResume(cleaned)
                        prediction = classify_resume(cleaned)
                        predicted.append(prediction)
                        extText = extract_text_from_docx(doc_file)
                        exp = expDetails(extText)
                        experience.append(exp)
                        skills.append(extract_skills(extText))
                        
                    elif file_extenstion == "Docx":
                         filename.append(doc_file.name)   
                         cleaned=clean_text(extract_text_from_docx(doc_file))
                         cleaned = cleanResume(cleaned)
                         prediction = classify_resume(cleaned)
                         predicted.append(prediction)
                         #predicted.append(classifier.predict(vectorizer1.transform([cleaned])))
                         extText = extract_text_from_docx(doc_file)
                         exp = expDetails(extText)
                         experience.append(exp)
                         skills.append(extract_skills(extText))    

        if len(predicted) > 0:
            file_type['Uploaded File'] = filename
            file_type['Experience'] = experience
            file_type['Skills'] = skills
            file_type['Predicted Profile'] = predicted
            # file_type
            st.table(file_type.style.format({'Experience': '{:.1f}'}))
            
            
            
        
        
    
if menu_id == 'Data Analysis':

    tab1, tab2, tab3 = st.tabs(["📁 Multiple Profiles and Text Distribution","🔍 Text Analysis","🔤 WordCloud"])
    df = pd.read_csv(r"E:\project2\ForDeployement.csv")
    workday = df[df["Resume"]=="workday"]
    react = df[df["Resume"]=="react developer"]
    peoplesoft = df[df["Resume"]=="peoplesoft"]
    sql = df[df["Resume"]=="sql developer"]
    workday_text_len = workday["cleantext"].str.len()
    react_text_len = react["cleantext"].str.len()
    peoplesoft_text_len = peoplesoft["cleantext"].str.len()
    sql_text_len = sql["cleantext"].str.len()
    workday_text = workday.cleantext.values
    react_text = react.cleantext.values
    peoplesoft_text = peoplesoft.cleantext.values
    sql_text = sql.cleantext.values
    def get_ngrams(text, ngram_from=2, ngram_to=2, n=None, max_features=20000):
    
        vec = CountVectorizer(ngram_range = (ngram_from, ngram_to), 
                                max_features = max_features, 
                                stop_words='english').fit(text)
        bag_of_words = vec.transform(text)
        sum_words = bag_of_words.sum(axis = 0) 
        words_freq = [(word, sum_words[0, i]) for word, i in vec.vocabulary_.items()]
        words_freq = sorted(words_freq, key = lambda x: x[1], reverse = True)
   
        return words_freq[:n]
    def generate_word_cloud(text):
        wordcloud = WordCloud(
                         width = 3000,
        height = 2000,
        background_color = 'White').generate(str(text))
    #    fig = plt.figure(
          #figsize = (5, 5))
        plt.imshow(wordcloud, interpolation = 'bilinear')
        plt.axis('off')
        plt.tight_layout(pad=0)
        plt.show()
    with tab1:

        col1, col2 = st.columns(2)

        with col1:
            st.subheader("Count of Resumes")
            plt.figure(figsize=(10, 6))
            
            sns.countplot(x = df["Resume"],order = df["Resume"].value_counts().index)
            st.pyplot()
        with col2:
            st.subheader("Distribution of Resume length")
            df['length'] = df['Text'].str.len()
            plt.figure(figsize=(10,6))
            sns.histplot(df['length'], kde=True,
            stat="density", kde_kws=dict(cut=3)).set_title('Resume length distribution')
            st.pyplot()
            
    
    with tab2:

        col1, col2, col3, col4, col5, col6, col7 = st.columns(7)

        with col1:
            combined_text = " ".join(df["Text"].tolist())
            tokens = combined_text.split()

            word_freq = Counter(tokens)

            # Get the most common words and their frequencies
            most_common_words = word_freq.most_common(20)
            words, frequencies = zip(*most_common_words)

            # Create a DataFrame from the most common words
            df_word_freq = pd.DataFrame({"Word": words, "Frequency": frequencies})
            plt.figure(figsize=(6,6))
            sns.barplot(data=df_word_freq, x="Frequency", y="Word")
            plt.xlabel("Frequency")
            plt.ylabel("Word")
            plt.title("Most Common Words")
            st.pyplot()
        
        with col2:
            combined_text = " ".join(df["cleantext"].tolist())
            tokens = combined_text.split()

            word_freq = Counter(tokens)

            # Get the most common words and their frequencies
            most_common_words = word_freq.most_common(20)
            words, frequencies = zip(*most_common_words)

            # Create a DataFrame from the most common words
            df_word_freq = pd.DataFrame({"Word": words, "Frequency": frequencies})

            sns.barplot(data=df_word_freq, x="Frequency", y="Word")
            plt.xlabel("Frequency")
            plt.ylabel("Word")
            plt.title("Most Common Words")
            st.pyplot()


           
        with col3:
            unigrams = get_ngrams(sql['cleantext'], ngram_from=1, ngram_to=1, n=15)
            unigrams_df = pd.DataFrame(unigrams)
            unigrams_df.columns=["Unigram", "Frequency"]
            sns.barplot(data=unigrams_df, x="Frequency", y="Unigram")
            plt.xlabel("Frequency")
            plt.ylabel("Word")
            plt.title("Most frequent Words in sql Resume")
            st.pyplot()

            
        with col4:
            unigrams = get_ngrams(react['cleantext'], ngram_from=1, ngram_to=1, n=15)
            unigrams_df = pd.DataFrame(unigrams)
            unigrams_df.columns=["Unigram", "Frequency"]
            sns.barplot(data=unigrams_df, x="Frequency", y="Unigram")
            plt.xlabel("Frequency")
            plt.ylabel("Word")
            plt.title("Most frequent Words in React Resume ")
            st.pyplot()

           
        with col5:
            unigrams = get_ngrams(peoplesoft['cleantext'], ngram_from=1, ngram_to=1, n=15)
            unigrams_df = pd.DataFrame(unigrams)
            unigrams_df.columns=["Unigram", "Frequency"]
            sns.barplot(data=unigrams_df, x="Frequency", y="Unigram")
            plt.xlabel("Frequency")
            plt.ylabel("Word")
            plt.title("Most frequent Words in peoplesoft Resume")
            st.pyplot()


           
        with col6:
            unigrams = get_ngrams(workday['cleantext'], ngram_from=1, ngram_to=1, n=15)
            unigrams_df = pd.DataFrame(unigrams)
            unigrams_df.columns=["Unigram", "Frequency"]
            sns.barplot(data=unigrams_df, x="Frequency", y="Unigram")
            plt.xlabel("Frequency")
            plt.ylabel("Word")
            plt.title("Most frequent Words in workday Resume")
            st.pyplot()

           
        with col7:
            unigrams = get_ngrams(df['cleantext'], ngram_from=2, ngram_to=2, n=15)
            unigrams_df = pd.DataFrame(unigrams)
            unigrams_df.columns=["Unigram", "Frequency"]
            sns.barplot(data=unigrams_df, x="Frequency", y="Unigram")
            plt.xlabel("Frequency")
            plt.ylabel("Word")
            plt.title("15 Most Common ngrams")
            st.pyplot()

            
    with tab3:
        
        ab1, ab2 = st.tabs(["🔤 WordCloud For non-CleanText and Clean Text","🔤 Profiles WordCloud "])

        with ab1:

            col1, col2 = st.columns(2)

            with col1:
                st.subheader("Word Cloud non-Clean Text")
                ttext = ' '.join(df["Text"].tolist())
                Q_wordcloud=WordCloud(
                    background_color='White',
                    width=2000,
                    height=2000
                   ).generate(ttext)
                fig = plt.figure(figsize = (5, 10))
                plt.axis('off')
                plt.imshow(Q_wordcloud)
                st.pyplot()

                
            with col2:
                st.subheader("Word Cloud Clean Text")
                ttext = ' '.join(df["cleantext"].tolist())
                Q_wordcloud=WordCloud(
                    background_color='White',
                    width=2000,
                    height=2000
                   ).generate(ttext)
                fig = plt.figure(figsize = (5, 10))
                plt.axis('off')
                plt.imshow(Q_wordcloud)
                st.pyplot()

                
        with ab2:

            col1, col2,col3,col4 = st.columns(4)

            with col1:
                st.subheader("Word Cloud of Workday Resume")
                generate_word_cloud(workday_text)
                st.pyplot()

                

            with col2:
                st.subheader("Word Cloud of React Developer Resume")
                generate_word_cloud(react_text)
                st.pyplot()
            with col3:
                st.subheader("Word Cloud of SQL Developer Resume")
                generate_word_cloud(sql_text)
                st.pyplot()
            with col4:
                st.subheader("Word Cloud of Peoplesoft Resume")
                generate_word_cloud(peoplesoft_text)
                st.pyplot()
                
if menu_id == "About":
    st.markdown("""<style>.stProgress .st-bo {color: purple;}</style>""", unsafe_allow_html=True)

    progress = st.progress(0)
    for i in range(100):
        progress.progress(i+1)
        sleep(0.001)

    st.markdown("<h1 style='text-align: center; '>BUSINESS OBJECTIVE </h1>", unsafe_allow_html=True)
    st.markdown("<h2 style='text-align: justify; font-size:180%; font-style: italic;'> Company’s usually receive huge number of Resumes/CV and have lots of financial documents everyday.The document classification solution should significantly reduce the manual human effort in the HRM and financial department. It should achieve a higher level of accuracy and automation with minimal human intervention.</h2>", unsafe_allow_html=True)
    st.markdown("<h2 style='text-align: center;'> ABSTRACT </h2>", unsafe_allow_html=True)
    st.markdown("<p style='text-align: justify; font-size:140%;'> A resume is a brief summary of your skills and experience. Companies recruiters and HR teams have a tough time scanning thousands of qualified resumes. Spending too many labor hours segregating candidates resume's manually is a waste of a company's time, money, and productivity. Recruiters, therefore, use resume classification in order to streamline the resume and applicant screening process. NLP technology allows recruiters to electronically gather, store, and organize large quantities of resumes. Once acquired, the resume data can be easily searched through and analyzed.Resumes are an ideal example of unstructured data. Since there is no widely accepted resume layout, each resume may have its own style of formatting, different text blocks and different category titles. Building a resume classification and gathering text from it is no easy task as there are so many kinds of layouts of resumes that you could imagine..</p>", unsafe_allow_html=True)
    st.markdown("<h2 style='text-align: center;'> INTRODUCTION </h2>", unsafe_allow_html=True)
    st.markdown("<p style='text-align: justify; font-size:140%; '>In this project we dive into building a Machine learning model for Resume Classification using Python and basic Natural language processing techniques. We would be using Python's libraries to implement various NLP (natural language processing) techniques like tokenization, lemmatization, parts of speech tagging, etc..</p>", unsafe_allow_html=True)
    st.markdown("<p style='text-align: justify; font-size:140%; '>A resume classification technology needs to be implemented in order to make it easy for the companies to process the huge number of resumes that are received by the organizations. This technology converts an unstructured form of resume data into a structured data format. The resumes received are in the form of documents from which the data needs to be extracted first such that the text can be classified or predicted based on the requirements. A resume classification analyzes resume data and extracts the information into the machine readable output. It helps automatically store, organize, and analyze the resume data to find out the candidate for the particular job position and requirements. This thus helps the organizations eliminate the error-prone and time-consuming process of going through thousands of resumes manually and aids in improving the recruiters’ efficiency.The basic data analysis process is performed such as data collection, data cleaning, exploratory data analysis, data visualization, and model building. The dataset consists of two columns, namely, Role Applied and Resume, where ‘role applied’ column is the domain field of the industry and ‘resume’ column consists of the text extracted from the resume document for each domain and industry.</p>", unsafe_allow_html=True)
    st.markdown("<p style='text-align: justify; font-size:140%; '>The aim of this project is achieved by performing the various data analytical methods and using the Machine Learning models and Natural Language Processing which will help in classifying the categories of the resume and building the Resume Classification Model..</p>", unsafe_allow_html=True)