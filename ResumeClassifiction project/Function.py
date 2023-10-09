import os #Used for path
import re #Regex
import csv #text into csv file
from docx import Document #.Docx Convert into text
from PyPDF2 import PdfReader #pdf Convert into text
import win32com.client as win32 #used for open Word file Component Object Model (COM)
from win32com.client import constants
import pandas as pd
import glob
import streamlit as st
import spacy 
nlp = spacy.load('en_core_web_md') #loading spacy pre built model
from spacy.matcher import Matcher 
from docx2pdf import convert
from io import BytesIO
import win32com.client as win32
from docx import Document



@st.cache_data
def extract_text_from_pdf(file_path): #This function is used for extract text from pdf
    pdf_reader = PdfReader(file_path)
    #num_pages = len(pdf_reader.pages)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text()
    return text


    #for page in range(file_path):
       # page_obj = pdf_reader.getPage(page)
        #text += page_obj.extractText()

   # return text
    
@st.cache_data
def extract_text_from_docx(file_path): #This function is used for extract text from Docx
    doc = Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs]
    return '\n'.join(paragraphs)

@st.cache_data
def save_as_docx(file_contents): #If Doc is not readed that this function convert Doc into Docx.by open Docx file and saveAS into Doc
    # Opening MS Word
    file_extension = os.path.splitext(file_contents.name)[1]

    if file_extension.lower() == '.docx':
        return extract_text_from_docx(file_contents)
    elif file_extension.lower() == '.doc':
        pdf_path = os.path.splitext(file_contents.name)[0] + '.pdf'
        convert(file_contents, pdf_path)

        docx_path = os.path.splitext(file_contents.name)[0] + '.docx'
        doc = Document()
        doc.add_paragraph(extract_text_from_pdf(pdf_path))
        doc.save(docx_path)

        os.remove(pdf_path)  # Remove the temporary PDF file
        return extract_text_from_docx(docx_path)
import textract
import tempfile
@st.cache_data


def extract_text_from_doc(file_path):
    # Open Word application
    word = win32.gencache.EnsureDispatch('Word.Application')
    doc = word.Documents.Open(file_path)
    doc.Activate()

    # Create a new file path for DOCX
    new_file_abs = os.path.splitext(file_path)[0] + '.docx'

    # Save as DOCX format
    word.ActiveDocument.SaveAs(new_file_abs, FileFormat=win32.constants.wdFormatXMLDocument)
    doc.Close(False)

    # Extract text from the converted DOCX file
    docx_doc = Document(new_file_abs)
    paragraphs = [p.text for p in docx_doc.paragraphs]
    text = '\n'.join(paragraphs)

    # Remove the temporary DOCX file
    os.remove(new_file_abs)

    return text


@st.cache_data
def clean_text(text):
    cleaned_text = re.sub(r'\s+', ' ', text)  # Remove extra space
    return cleaned_text


@st.cache_data
def extract_skills(resume_text):
    nlp_text = nlp(resume_text)

    # removing stop words and implementing word tokenization
    tokens = [token.text for token in nlp_text if not token.is_stop]
    # Set of skills if skills match with this text than they append in skills column
    skills = ["c","my sql","nodejs","node.js","reactjs","javascript","html","css","javascript","angular js",
              "js","fcsm","sql developer","peopleSoft","mysql","sql","plsql","nosql","rdbms","ddl" ,"dml","dcl"
              ,"sql developer","core hcm","xml", "xslt", "eib", "core connectors","workday",'hcm',"peoplesoft admin",
               "PeopleSoft Admin","dba","" ]
    
    skillset = []
    
    # check for one-grams (example: python)
    for token in tokens:
        if token.lower() in skills:
            skillset.append(token)
    
    #check for bi-grams and tri-grams (example: machine learning)
    for token in nlp_text.noun_chunks:
        token = token.text.lower().strip()
        if token in skills:
            skillset.append(token)
    
  
    return ', '.join(set([i.capitalize() for i in skillset]))



# In name extraction we use POS tagging. 
matcher = Matcher(nlp.vocab) 
@st.cache_data
def extract_name(resume_text):
    nlp_text = nlp(resume_text)
    
    # First name and Last name are always Proper Nouns
    pattern = [{'POS': 'PROPN'}, {'POS': 'PROPN'}]
    
    matcher.add('NAME', [pattern], on_match = None)
    
    matches = matcher(nlp_text)
    
    for match_id, start, end in matches:
        span = nlp_text[start:end]
        return span.text


nlp = spacy.load('en_core_web_lg')
@st.cache_data
def extract_experiences(text): #Extract Experience from resume
    doc = nlp(text)

    experiences = []
    for sent in doc.sents:
        if 'experience' in sent.text.lower():
            experiences.append(sent.text)

    return ', '.join(set(i for i in experiences))



@st.cache_data
def ed(text):  #Extract Education from Resume
    doc = nlp(text)

    edu = []
    for sent in doc.sents:
        if ("educational qualification" and 'education') in sent.text.lower():
            edu.append(sent.text)
    return ', '.join(set(i for i in edu))



import re
@st.cache_data
def extract_website_links(text):
    # Regular expression pattern to match website links
    pattern = r'(https?://\S+)'

    # Find all matches using the pattern
    matches = re.findall(pattern, text)

    return ", ".join(matches)


@st.cache_data
def extract_phone_numbers(text):
    # Regular expression pattern to match phone numbers
    pattern = r'(?<!\d)(?:(?:\d{2}[-\s]?\d{8})|(?:\d{4}[-\s]?\d{6}))(?!\d)'

    # Find all matches using the pattern
    matches = re.findall(pattern, text)

    return ", ".join(matches)


@st.cache_data
def expDetails(Text):  #Extract experience in year from Resumes
    global sent
   
    Text = Text.split()
   
    for i in range(len(Text)-2):
        Text[i].lower()
        
        if Text[i] ==  'years':
            sent =  Text[i-2] + ' ' + Text[i-1] +' ' + Text[i] +' '+ Text[i+1] +' ' + Text[i+2]
            l = re.findall('\d*\.?\d+',sent)
            for i in l:
                a = float(i)
            return(a)
            return (sent)


def cleanResume(resumeText):
    if pd.isnull(resumeText) or not isinstance(resumeText, str):
        return ''
    #resumeText = re.sub(r'\d+', '', resumeText)
    resumeText = re.sub('•', '',resumeText)
    resumeText = re.sub(r'\s+', ' ', resumeText).strip()
    resumeText = re.sub('http\S+\s*', ' ', resumeText)  # remove URLs
    resumeText = re.sub('RT|cc', ' ', resumeText)  # remove RT and cc
    resumeText = re.sub('#\S+', '', resumeText)  # remove hashtags
    resumeText = re.sub('@\S+', '  ', resumeText)  # remove mentions
    resumeText = re.sub('[%s]' % re.escape("""!"#$%&'()*+,-./:;<=>?@[\]^_`{|}~"""), ' ', resumeText)  # remove punctuations
    resumeText = re.sub(r'[^\x00-\x7f]',r' ', resumeText) 
    resumeText = re.sub('\s+', ' ', resumeText)  # remove extra whitespace
    return resumeText        

#input_directory = r'C:\Users\abc\Desktop\Resume Classification\Resume1'
#output_csv = 'Resume_text.csv'

#with open(output_csv, 'w', newline='', encoding='utf-8') as csv_file:
    #writer = csv.writer(csv_file)
    #writer.writerow(['Text',"Resume"])  # Write header row
    #for file_path in glob.glob(os.path.join(input_directory, '**', '*'), recursive=True):
        #if not os.path.isfile(file_path):
            #continue

  